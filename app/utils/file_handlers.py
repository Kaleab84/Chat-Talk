from pathlib import Path
from typing import Dict, Any, Optional
import logging
from docx import Document

logger = logging.getLogger(__name__)

class FileHandler:
    """Handles different file types for document processing."""
    
    @staticmethod
    def read_docx(file_path: Path) -> str:
        """Extract text from .docx file."""
        try:
            doc = Document(str(file_path))
            paragraphs = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            return "\n".join(paragraphs)
            
        except Exception as e:
            logger.error(f"Error reading .docx file {file_path}: {e}")
            raise
    
    @staticmethod
    def read_doc(file_path: Path) -> str:
        """Extract text from .doc file using COM (Windows only)."""
        try:
            import win32com.client
            
            word = win32com.client.Dispatch("Word.Application")
            word.visible = False
            doc = word.Documents.Open(str(file_path))
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            
            return text
            
        except ImportError:
            logger.error("win32com.client not available. Cannot read .doc files.")
            raise ValueError("Cannot process .doc files without win32com.client")
        except Exception as e:
            logger.error(f"Error reading .doc file {file_path}: {e}")
            raise
    
    @staticmethod
    def read_txt(file_path: Path, encoding: str = 'utf-8') -> str:
        """Extract text from .txt file."""
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            # Try different encodings
            for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_path.read_text(encoding=enc)
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode file {file_path} with any encoding")
        except Exception as e:
            logger.error(f"Error reading .txt file {file_path}: {e}")
            raise
    
    @staticmethod
    def get_file_info(file_path: Path) -> Dict[str, Any]:
        """Get basic file information."""
        try:
            stat = file_path.stat()
            return {
                'name': file_path.name,
                'size': stat.st_size,
                'modified': stat.st_mtime,
                'extension': file_path.suffix.lower(),
                'stem': file_path.stem
            }
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            return {}
    
    @staticmethod
    def is_supported_format(file_path: Path) -> bool:
        """Check if file format is supported."""
        supported_extensions = {'.docx', '.doc', '.txt'}
        return file_path.suffix.lower() in supported_extensions